"""Tests for Word (DOCX) to PDF conversion service."""

from pathlib import Path

import pytest
from docx import Document

from backend.services.word_service import word_to_pdf


@pytest.fixture()
def docx_file(tmp_path: Path) -> Path:
    """Create a sample DOCX file with headings and paragraphs."""
    doc = Document()
    doc.add_heading("Hello World", level=1)
    doc.add_paragraph("This is a bold and italic paragraph.")
    doc.add_heading("Lists", level=2)
    doc.add_paragraph("First item", style="List Bullet")
    doc.add_paragraph("Second item", style="List Bullet")
    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


@pytest.fixture()
def docx_with_table(tmp_path: Path) -> Path:
    """Create a DOCX file with a table."""
    doc = Document()
    doc.add_heading("Report", level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "100"
    table.cell(2, 0).text = "B"
    table.cell(2, 1).text = "200"
    path = tmp_path / "table.docx"
    doc.save(str(path))
    return path


class TestWordToPdf:
    """Tests for word_to_pdf."""

    def test_basic_conversion(self, docx_file: Path, tmp_path: Path):
        output = tmp_path / "output.pdf"
        result = word_to_pdf(docx_file, output)

        assert result == output
        assert output.exists()
        assert output.stat().st_size > 0
        assert output.read_bytes()[:5] == b"%PDF-"

    def test_table_conversion(self, docx_with_table: Path, tmp_path: Path):
        output = tmp_path / "output.pdf"
        result = word_to_pdf(docx_with_table, output)

        assert result == output
        assert output.exists()
        assert output.read_bytes()[:5] == b"%PDF-"

    def test_letter_paper_size(self, docx_file: Path, tmp_path: Path):
        output = tmp_path / "output.pdf"
        result = word_to_pdf(docx_file, output, paper_size="letter")

        assert result == output
        assert output.exists()
        assert output.read_bytes()[:5] == b"%PDF-"

    def test_empty_docx(self, tmp_path: Path):
        doc = Document()
        path = tmp_path / "empty.docx"
        doc.save(str(path))
        output = tmp_path / "output.pdf"

        result = word_to_pdf(path, output)
        assert result == output
        assert output.exists()
